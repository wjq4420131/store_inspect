# 作者：wjq   
# 创建时间: 2019/5/27  
# 文件: test30.py   
# 软件名称: PyCharm

from docx import Document

document = Document()
document.add_heading('巡检表', 0)
document.add_heading('数据库服务器', level=1)
document.add_heading('中心', level=1)

p = document.add_heading[0]
print(p)