import paramiko
import datetime
import os
import time
from docx import Document


# 读取的脚本功能 1:巡检内容  2:负载状态


while True:
    print("1、数据库","2、中心","3、接入","4、分发","5、sip","6、上级媒体","0、表示退出",)
    #print("每次只能输入一项，以数字代替")
    #func = int(input("请输入巡检内容:"))
    #func = input("请输入巡检内容:")
    func = list(map(int, input('请输入巡检内容以逗号分隔,逗号必须是英文状态下的:').split(',')))
    for func in func:
        print(int(func))

        # 读取当前路径
        base_dir = os.getcwd()
        # 命令开始执行时间
        starttime = datetime.datetime.now()
        print(" -------------------------------------------------------------")
        print("|                                                             |")
        print("  startime:        ", starttime)
        print("|                                                             |")
        print(" -------------------------------------------------------------")
        # 注意路径前面的r，否则有些文件会当作转义字符处理
        # 读取命令脚本
        if func == 1:
            cmd_filepath = base_dir + r"\cmd\shujuku.txt"
            ip_filepath = base_dir + r"\server_ip\shujuku.txt"
        if func == 2:
            cmd_filepath = base_dir + r"\cmd\zhongxin.txt"
            ip_filepath = base_dir + r"\server_ip\zhongxin.txt"
        if func == 3:
            cmd_filepath = base_dir + r"\cmd\jieru.txt"
            ip_filepath = base_dir + r"\server_ip\jieru.txt"
        if func == 4:
            cmd_filepath = base_dir + r"\cmd\fenfa.txt"
            ip_filepath = base_dir + r"\server_ip\fenfa.txt"
        if func == 5:
            cmd_filepath = base_dir + r"\cmd\sip.txt"
            ip_filepath = base_dir + r"\server_ip\sip.txt"
        if func == 6:
            cmd_filepath = base_dir + r"\cmd\meiti.txt"
            ip_filepath = base_dir + r"\server_ip\meiti.txt"
        cmd_file = open(cmd_filepath, "r",encoding='gbk')
        cmd = cmd_file.read()
        #print(cmd)

    #,encoding='gbk'
        # 按日期创建巡检结果
        document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
        document.add_heading('巡检表', 0)
        document.add_paragraph('巡检人：')
        document.add_paragraph('巡检时间：')
        document.add_paragraph('巡检中发现的问题：')
        table = document.add_table(rows=4, cols=6,style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '业务模块'
        hdr_cells[1].text = 'IP地址'
        hdr_cells[2].text = '问题描述'
        hdr_cells[3].text = '处理过程'
        hdr_cells[4].text = '遗留问题'
        hdr_cells[5].text = '处理人'
        dayTime = datetime.datetime.now().strftime('%Y-%m-%d')
        pwd_file = os.getcwd() + '\\' + 'result' + '\\' + dayTime + '\\'
        isExists = os.path.exists(pwd_file)
        if not isExists:
            os.makedirs(pwd_file)

        # print(cmd)
        # 读取IP地址列表
        ip_file = open(ip_filepath, "r")
        while 1:
            ipinfo = ip_file.readline()
            if not ipinfo:
                break
            else:
                # 读取IP，用户名，密码
                infos = ipinfo.split(',')
                host = infos[0]
                port = infos[1]
                username = infos[2]
                pwd = infos[3].strip()
                pwd = pwd.strip('\n')
                # 远程连接服务器
            try:
                client = paramiko.SSHClient()
                client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
                client.connect(host, port, username, pwd, timeout=40)
                stdin, stdout, stderr = client.exec_command(cmd)

                endtime = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                if func == 1:
                    document.add_heading('oracle数据库：' + host)
                    p = document.add_paragraph()
                    for line in stdout:
                        p.add_run(line.strip('\n') + '\n')
                        print(line.strip('\n'))
                    paragraph = document.add_paragraph(endtime,style='List Bullet')
                    document.save(pwd_file + '\oracle数据库.doc')
                    #
                    # with open(pwd_file + '\shujuku.txt', 'a+') as file_txt:
                    #     file_txt.write(
                    #         "===============================================================================================================" + '\n')
                    #     file_txt.write("                                                   " + host + "      " + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    #     file_txt.write("" + '\n')
                    #     for line in stdout:
                    #         file_txt.write(line.strip('\n') + '\n')
                    #         print(line.strip('\n'))
                    #     print(endtime)
                    #     file_txt.write(
                    #         "===============================================================================================================" + '\n')
                    #     file_txt.write(                                                  endtime + '\n' )
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    # file_txt.close()
                    # # print(line.strip('\n'))
                    # client.close()
                    # print("")
                if func == 2:
                    document.add_heading('中心服务器：' + host)
                    p = document.add_paragraph()
                    for line in stdout:
                        p.add_run(line.strip('\n') + '\n')
                        print(line.strip('\n'))
                    paragraph = document.add_paragraph(endtime,style='List Bullet')
                    document.save(pwd_file + '\中心.doc')
                    # with open(pwd_file + '\zhongxin.txt', 'a+') as file_txt:
                    #     file_txt.write(
                    #         "==================================================================================================================" + '\n')
                    #     file_txt.write("                                                      " + host + "      " + '\n')
                    #     file_txt.write(
                    #         "===================================================================================================================" + '\n')
                    #     file_txt.write("" + '\n')
                    #     for line in stdout:
                    #         file_txt.write(line.strip('\n') + '\n')
                    #         print(line.strip('\n'))
                    #     print(endtime)
                    #     file_txt.write(
                    #         "===============================================================================================================" + '\n')
                    #     file_txt.write(                                                  endtime + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    #
                    # file_txt.close()
                    # # print(line.strip('\n'))
                    # client.close()
                    # print("")
                if func == 3:
                    document.add_heading('接入服务器：' + host)
                    p = document.add_paragraph()
                    for line in stdout:
                        p.add_run(line.strip('\n') + '\n')
                        print(line.strip('\n'))
                    paragraph = document.add_paragraph(endtime,style='List Bullet')
                    document.save(pwd_file + '\接入.doc')
                    # with open(pwd_file + '\jieru.txt', 'a+') as file_txt:
                    #     file_txt.write(
                    #         "===================================================================================================================" + '\n')
                    #     file_txt.write("                                                         " + host + "      " + '\n')
                    #     file_txt.write(
                    #         "====================================================================================================================" + '\n')
                    #     file_txt.write("" + '\n')
                    #     for line in stdout:
                    #         file_txt.write(line.strip('\n') + '\n')
                    #         print(line.strip('\n'))
                    #     print(endtime)
                    #     file_txt.write(
                    #         "===============================================================================================================" + '\n')
                    #     file_txt.write(                                                  endtime + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    # file_txt.close()
                    # # print(line.strip('\n'))
                    # client.close()
                    # print("")
                if func == 4:
                    document.add_heading('分发服务器：' + host)
                    p = document.add_paragraph()
                    for line in stdout:
                        p.add_run(line.strip('\n') + '\n')
                        print(line.strip('\n'))
                    paragraph = document.add_paragraph(endtime,style='List Bullet')
                    document.save(pwd_file + '\分发.doc')
                    # with open(pwd_file + r'\fenfa.txt', 'a+') as file_txt:
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    #     file_txt.write("                                                   " + host + "      " + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    #     file_txt.write("" + '\n')
                    #     for line in stdout:
                    #         file_txt.write(line.strip('\n') + '\n')
                    #         print(line.strip('\n'))
                    #     print(endtime)
                    #     file_txt.write(
                    #         "===============================================================================================================" + '\n')
                    #     file_txt.write(                                                  endtime + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    # file_txt.close()
                    # # print(line.strip('\n'))
                    # client.close()
                    # print("")
                if func == 5:
                    document.add_heading('Sip服务器：' + host)
                    p = document.add_paragraph()
                    for line in stdout:
                        p.add_run(line.strip('\n') + '\n')
                        print(line.strip('\n'))
                    paragraph = document.add_paragraph(endtime,style='List Bullet')
                    document.save(pwd_file + '\sip服务器.doc')
                    # with open(pwd_file + '\sip.txt', 'a+') as file_txt:
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    #     file_txt.write("                                                   " + host + "      " + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    #     file_txt.write("" + '\n')
                    #     for line in stdout:
                    #         file_txt.write(line.strip('\n') + '\n')
                    #         print(line.strip('\n'))
                    #     print(endtime)
                    #     file_txt.write(
                    #         "===============================================================================================================" + '\n')
                    #     file_txt.write(                                                  endtime + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    # file_txt.close()
                    # # print(line.strip('\n'))
                    # client.close()
                    # print("")
                if func == 6:
                    # with open(pwd_file + '\meiti.txt', 'a+') as file_txt:
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    #     file_txt.write("                                                   " + host + "      " + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    #     file_txt.write("" + '\n')
                    #     for line in stdout:
                    #         file_txt.write(line.strip('\n') + '\n')
                    #         print(line.strip('\n'))
                    #     print(endtime)
                    #     file_txt.write(
                    #         "===============================================================================================================" + '\n')
                    #     file_txt.write(                                                  endtime + '\n')
                    #     file_txt.write(
                    #         "================================================================================================================" + '\n')
                    # file_txt.close()
                    # # print(line.strip('\n'))
                    # client.close()
                    # print("")
                    document.add_heading('上级媒体服务器：' + host)
                    p = document.add_paragraph()
                    for line in stdout:
                        p.add_run(line.strip('\n') + '\n')
                        print(line.strip('\n'))
                    paragraph = document.add_paragraph(endtime,style='List Bullet')
                    document.save(pwd_file + '\上级媒体.doc')
            except Exception as e:
                print(host)
                with open(pwd_file + '\error.txt', 'a+') as error_file:
                    error_file.write(host + '\n')
                    error_file.close()
                #raise e

        # 命令执行完成时间

        endtime = datetime.datetime.now()
        print(" -------------------------------------------------------------")
        print("|                                                             |")
        print("    endtime:      ", endtime)
        print("|                                                             |")
        print(" -------------------------------------------------------------")
        print(" -------------------------------------------------------------")
        print("|                                                             |")
        print("  startime:       ", starttime)
        print("  endtime:        ", endtime)
        print("  cost:           ", endtime - starttime)
        print("|                                                             |")
        print(" -------------------------------------------------------------")
        if func==0:
            break
